"""
Output formatters: Markdown, HTML, PDF (reportlab), DOCX (python-docx).
"""
import os
from datetime import datetime


# ---------------------------------------------------------------------------
# Markdown
# ---------------------------------------------------------------------------

def _to_markdown(items: list[dict], action: str) -> str:
    lines = []
    lines.append(f"# Transcript Export\n")
    lines.append(f"*Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}*\n")

    for item in items:
        lines.append(f"\n---\n\n## {item['title']}\n")
        if item.get("video_url"):
            lines.append(f"*Source: {item['video_url']}*\n")

        if action == "plain":
            lines.append(f"\n{item['body']}\n")

        elif action == "summary":
            orig = item.get("word_count_original", "?")
            summ = item.get("word_count_summary", "?")
            lines.append(f"\n*Summary ({summ} words, down from {orig})*\n")
            lines.append(f"\n{item['body']}\n")

        elif action == "booklet":
            for para in item.get("paragraphs", []):
                lines.append(f"\n{para}\n")

        elif action == "magazine":
            for sec in item.get("sections", []):
                lines.append(f"\n### {sec['heading']}\n")
                if sec.get("image"):
                    img = sec["image"]
                    lines.append(f"\n![{img['caption']}]({img['url']})\n")
                    lines.append(f"*{img['caption']}*\n")
                lines.append(f"\n{sec['body']}\n")

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------

_HTML_BASE = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Transcript Export</title>
<style>
  body {{ font-family: Georgia, serif; max-width: 820px; margin: 40px auto; padding: 0 24px;
          color: #222; line-height: 1.75; background: #fafaf8; }}
  h1 {{ font-size: 2em; border-bottom: 2px solid #ddd; padding-bottom: 12px; }}
  h2 {{ font-size: 1.5em; margin-top: 2.5em; color: #1a1a1a; }}
  h3 {{ font-size: 1.15em; color: #444; margin-top: 1.8em; }}
  .meta {{ color: #888; font-size: 0.9em; font-style: italic; }}
  .section {{ margin-bottom: 2em; }}
  .magazine-section {{ display: grid; grid-template-columns: 1fr; gap: 1em;
                        border-top: 1px solid #e0e0e0; padding-top: 1.5em; }}
  .magazine-section.has-image {{ grid-template-columns: 1fr 280px; }}
  .mag-image {{ text-align: center; }}
  .mag-image img {{ max-width: 100%; border-radius: 6px; }}
  .mag-image figcaption {{ font-size: 0.8em; color: #888; margin-top: 6px; }}
  .summary-stats {{ background: #f0f0ec; border-left: 3px solid #aaa;
                    padding: 8px 14px; font-size: 0.9em; color: #555; margin-bottom: 1em; }}
  a {{ color: #1a5276; }}
  @media (max-width: 600px) {{ .magazine-section.has-image {{ grid-template-columns: 1fr; }} }}
</style>
</head>
<body>
<h1>Transcript Export</h1>
<p class="meta">Generated {date}</p>
{body}
</body>
</html>"""


def _to_html(items: list[dict], action: str) -> str:
    body_parts = []
    for item in items:
        title = item["title"]
        url = item.get("video_url", "")
        src = f'<p class="meta">Source: <a href="{url}">{url}</a></p>' if url else ""
        body_parts.append(f'<hr>\n<h2>{title}</h2>\n{src}')

        if action == "plain":
            body_parts.append(f'<div class="section"><p>{item["body"]}</p></div>')

        elif action == "summary":
            orig = item.get("word_count_original", "?")
            summ = item.get("word_count_summary", "?")
            body_parts.append(
                f'<div class="summary-stats">Summary: {summ} words (condensed from {orig})</div>'
                f'<div class="section"><p>{item["body"]}</p></div>'
            )

        elif action == "booklet":
            paras = "".join(f"<p>{p}</p>" for p in item.get("paragraphs", []))
            body_parts.append(f'<div class="section">{paras}</div>')

        elif action == "magazine":
            for sec in item.get("sections", []):
                img_html = ""
                has_img = sec.get("image") is not None
                if has_img:
                    img = sec["image"]
                    img_html = (
                        f'<figure class="mag-image">'
                        f'<img src="{img["url"]}" alt="{img["caption"]}" loading="lazy">'
                        f'<figcaption>{img["caption"]}</figcaption></figure>'
                    )
                cls = "magazine-section has-image" if has_img else "magazine-section"
                body_parts.append(
                    f'<div class="{cls}">'
                    f'<div><h3>{sec["heading"]}</h3><p>{sec["body"]}</p></div>'
                    f'{img_html}</div>'
                )

    body = "\n".join(body_parts)
    return _HTML_BASE.format(date=datetime.now().strftime("%Y-%m-%d %H:%M"), body=body)


# ---------------------------------------------------------------------------
# PDF (reportlab)
# ---------------------------------------------------------------------------

def _to_pdf(items: list[dict], action: str, out_path: str):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    from reportlab.lib import colors

    doc = SimpleDocTemplate(out_path, pagesize=A4,
                            leftMargin=2.5*cm, rightMargin=2.5*cm,
                            topMargin=2.5*cm, bottomMargin=2.5*cm)
    styles = getSampleStyleSheet()
    style_body = ParagraphStyle("body", parent=styles["Normal"],
                                fontSize=11, leading=17, spaceAfter=10)
    style_h1 = ParagraphStyle("h1", parent=styles["Heading1"],
                               fontSize=20, spaceAfter=6)
    style_h2 = ParagraphStyle("h2", parent=styles["Heading2"],
                               fontSize=15, spaceBefore=24, spaceAfter=4)
    style_h3 = ParagraphStyle("h3", parent=styles["Heading3"],
                               fontSize=12, spaceBefore=14, spaceAfter=4)
    style_meta = ParagraphStyle("meta", parent=styles["Normal"],
                                fontSize=9, textColor=colors.grey, spaceAfter=6)
    style_note = ParagraphStyle("note", parent=styles["Normal"],
                                fontSize=9, textColor=colors.HexColor("#555555"),
                                backColor=colors.HexColor("#f0f0ec"),
                                leftIndent=10, rightIndent=10,
                                spaceAfter=10, leading=14)

    story = [
        Paragraph("Transcript Export", style_h1),
        Paragraph(f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}", style_meta),
    ]

    for item in items:
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.grey))
        story.append(Spacer(1, 0.3*cm))
        story.append(Paragraph(item["title"], style_h2))
        if item.get("video_url"):
            story.append(Paragraph(item["video_url"], style_meta))

        if action == "plain":
            story.append(Paragraph(item["body"], style_body))

        elif action == "summary":
            orig = item.get("word_count_original", "?")
            summ = item.get("word_count_summary", "?")
            story.append(Paragraph(f"Summary: {summ} words (from {orig})", style_note))
            story.append(Paragraph(item["body"], style_body))

        elif action == "booklet":
            for para in item.get("paragraphs", []):
                story.append(Paragraph(para, style_body))
                story.append(Spacer(1, 0.2*cm))

        elif action == "magazine":
            for sec in item.get("sections", []):
                story.append(Paragraph(sec["heading"], style_h3))
                if sec.get("image"):
                    # Attempt to embed image via URL
                    try:
                        import urllib.request
                        import tempfile
                        from reportlab.platypus import Image as RLImage
                        img_data = sec["image"]
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp:
                            urllib.request.urlretrieve(img_data["url"], tmp.name)
                            img_flow = RLImage(tmp.name, width=10*cm, height=6*cm)
                            img_flow.hAlign = "LEFT"
                            story.append(img_flow)
                            story.append(Paragraph(img_data["caption"], style_meta))
                    except Exception:
                        pass  # Skip image silently if it can't be fetched
                story.append(Paragraph(sec["body"], style_body))

    doc.build(story)


# ---------------------------------------------------------------------------
# DOCX (python-docx)
# ---------------------------------------------------------------------------

def _to_docx(items: list[dict], action: str, out_path: str):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    title_para = doc.add_heading("Transcript Export", 0)
    meta = doc.add_paragraph(f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    meta.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    meta.runs[0].font.size = Pt(9)

    for item in items:
        doc.add_paragraph("─" * 60)
        doc.add_heading(item["title"], level=1)
        if item.get("video_url"):
            url_p = doc.add_paragraph(item["video_url"])
            url_p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            url_p.runs[0].font.size = Pt(9)

        if action == "plain":
            doc.add_paragraph(item["body"])

        elif action == "summary":
            orig = item.get("word_count_original", "?")
            summ = item.get("word_count_summary", "?")
            note = doc.add_paragraph(f"Summary: {summ} words (condensed from {orig})")
            note.runs[0].font.size = Pt(9)
            note.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            doc.add_paragraph(item["body"])

        elif action == "booklet":
            for para in item.get("paragraphs", []):
                doc.add_paragraph(para)

        elif action == "magazine":
            for sec in item.get("sections", []):
                doc.add_heading(sec["heading"], level=2)
                if sec.get("image"):
                    try:
                        import urllib.request, tempfile
                        img_data = sec["image"]
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp:
                            urllib.request.urlretrieve(img_data["url"], tmp.name)
                            doc.add_picture(tmp.name, width=Cm(10))
                            cap = doc.add_paragraph(img_data["caption"])
                            cap.runs[0].font.size = Pt(8)
                            cap.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                    except Exception:
                        pass
                doc.add_paragraph(sec["body"])

    doc.save(out_path)


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def format_output(items: list[dict], fmt: str, out_path: str, action: str):
    if fmt == "md":
        content = _to_markdown(items, action)
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(content)

    elif fmt == "html":
        content = _to_html(items, action)
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(content)

    elif fmt == "pdf":
        _to_pdf(items, action, out_path)

    elif fmt == "docx":
        _to_docx(items, action, out_path)

    else:
        raise ValueError(f"Unknown format: {fmt}")
